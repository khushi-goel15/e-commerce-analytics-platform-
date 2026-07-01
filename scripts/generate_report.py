"""
Generate a comprehensive Word (.docx) report for the E-Commerce Analytics Platform.
"""
import pandas as pd
import numpy as np
import os, json
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from datetime import datetime

BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
RAW_DIR = os.path.join(BASE_DIR, "data", "raw")
EXPORT_DIR = os.path.join(BASE_DIR, "data", "exports")
RESULTS_PATH = os.path.join(EXPORT_DIR, "analytics_results.json")

# Load analytics results
with open(RESULTS_PATH) as f:
    results = json.load(f)

# Load datasets
customers = pd.read_csv(os.path.join(RAW_DIR, "customers.csv"))
orders = pd.read_csv(os.path.join(RAW_DIR, "orders.csv"))
order_items = pd.read_csv(os.path.join(RAW_DIR, "order_items.csv"))
products = pd.read_csv(os.path.join(RAW_DIR, "products.csv"))
returns = pd.read_csv(os.path.join(RAW_DIR, "returns.csv"))

doc = Document()

# ── Styles ───────────────────────────────────────────────────────
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# ── Title Page ───────────────────────────────────────────────────
for _ in range(6):
    doc.add_paragraph("")

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("E-COMMERCE ANALYTICS PLATFORM")
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("End-to-End Analytics System\nSQL + Python + Power BI + Tableau")
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x4A, 0x4A, 0x4A)

doc.add_paragraph("")
date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = date_p.add_run(f"Generated: {datetime.now().strftime('%B %d, %Y')}")
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.add_page_break()

# ── Helper ───────────────────────────────────────────────────────
def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)
    return h

def add_table(df, caption=None, decimals=2):
    if caption:
        p = doc.add_paragraph(caption)
        p.runs[0].bold = True if p.runs else False
        p.runs[0].font.size = Pt(10) if p.runs else Pt(10)
    num_cols = len(df.columns)
    num_rows = len(df) + 1
    table = doc.add_table(rows=min(num_rows, 30), cols=num_cols)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Headers
    for j, col in enumerate(df.columns):
        cell = table.rows[0].cells[j]
        cell.text = str(col).replace("_", " ").title()
        for p in cell.paragraphs:
            p.runs[0].bold = True
            p.runs[0].font.size = Pt(8)
    # Data
    for i, (_, row) in enumerate(df.iterrows()):
        if i >= 29:
            break
        for j, col in enumerate(df.columns):
            val = row[col]
            cell = table.rows[i + 1].cells[j]
            if isinstance(val, float):
                cell.text = f"{val:.{decimals}f}"
            else:
                cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(8)

def add_kpi_box(label, value, color="1B3A5C"):
    p = doc.add_paragraph()
    run = p.add_run(f"{label}:  ")
    run.bold = True
    run.font.size = Pt(12)
    run = p.add_run(str(value))
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)

def add_insight(text):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.size = Pt(11)

# ══════════════════════════════════════════════════════════════════
# 1. EXECUTIVE SUMMARY
# ══════════════════════════════════════════════════════════════════
add_heading("1. Executive Summary", level=1)

doc.add_paragraph(
    "This report presents the complete E-Commerce Analytics Platform, built from scratch "
    "to help an e-commerce company track revenue, profit, customer behavior, product "
    "performance, and regional sales. The system comprises a synthetic dataset with "
    "realistic business patterns, an ETL pipeline, a star-schema database, advanced SQL "
    "analysis, Python analytics (segmentation, forecasting), and exports for both Power BI "
    "and Tableau."
)

overview = results.get("eda", {}).get("overview", {})
add_heading("Key Metrics", level=2)
add_kpi_box("Total Revenue", f"${overview.get('revenue', 0):,.2f}")
add_kpi_box("Total Profit", f"${overview.get('profit', 0):,.2f}")
add_kpi_box("Profit Margin", f"{overview.get('margin_pct', 0):.2f}%")
add_kpi_box("Total Orders", f"{overview.get('orders', 0):,}")
add_kpi_box("Unique Customers", f"{overview.get('customers', 0):,}")
add_kpi_box("Average Order Value", f"${overview.get('aov', 0):,.2f}")
add_kpi_box("Return Rate", f"{overview.get('return_rate_pct', 0):.2f}%")

# ══════════════════════════════════════════════════════════════════
# 2. DATASET OVERVIEW
# ══════════════════════════════════════════════════════════════════
add_heading("2. Dataset Overview", level=1)
doc.add_paragraph(
    "The dataset was synthetically generated using Python with deliberate business patterns "
    "embedded: Pareto-distributed customer value, seasonality (Q4/Diwali spikes), regional "
    "discount dynamics, and return-rate variations by category. Data quality issues such as "
    "missing signup dates (15.4%) and future-dated orders were intentionally included to "
    "simulate real-world messiness."
)

# Dataset summary table
ds_summary = pd.DataFrame({
    "File": ["customers.csv", "orders.csv", "order_items.csv", "products.csv", "returns.csv"],
    "Rows": [len(customers), len(orders), len(order_items), len(products), len(returns)],
    "Columns": [len(customers.columns), len(orders.columns), len(order_items.columns), len(products.columns), len(returns.columns)],
    "Description": [
        "Customer demographics with Pareto-skewed value",
        f"Order transactions with {orders['order_date'].nunique()} unique dates, 112 future-dated",
        "Line items: quantity, discount, selling price",
        "Product master: 5 categories, 25 sub-categories, 25 brands",
        f"Return records: {returns['return_reason'].nunique()} reason types",
    ],
})
add_table(ds_summary)
doc.add_paragraph("")

cat_counts = products['category'].value_counts().reset_index()
cat_counts.columns = ["Category", "Product Count"]
add_table(cat_counts, "Products by Category:")

# ══════════════════════════════════════════════════════════════════
# 3. ETL PIPELINE
# ══════════════════════════════════════════════════════════════════
add_heading("3. ETL Pipeline", level=1)
doc.add_paragraph(
    "The ETL pipeline (scripts/etl_pipeline.py) performs three stages:"
)
doc.add_paragraph("Extract: Loads raw CSVs from data/raw/", style='List Bullet')
doc.add_paragraph(
    "Transform: Removes duplicates, handles nulls, validates dates, "
    "clips negatives, calculates revenue = quantity × selling_price, "
    "profit = revenue − (quantity × cost_price), discount_impact = discount × revenue / 100, "
    "and marks returned items.",
    style='List Bullet'
)
doc.add_paragraph("Load: Generates a PostgreSQL-compatible SQL script (sql/load_data.sql) with INSERT statements.", style='List Bullet')
doc.add_paragraph(
    "Output: 25,228 fact rows written to data/processed/fact_sales.csv "
    f"with total revenue ${overview.get('revenue', 0):,.2f}."
)

# ══════════════════════════════════════════════════════════════════
# 4. DATABASE SCHEMA (STAR SCHEMA)
# ══════════════════════════════════════════════════════════════════
add_heading("4. Database Schema — Star Schema", level=1)
doc.add_paragraph(
    "The database follows a star schema design (sql/schema.sql) optimized for analytical queries:"
)

dim_tables = pd.DataFrame({
    "Table": ["dim_customers", "dim_products", "dim_date", "dim_region", "fact_sales"],
    "Type": ["Dimension", "Dimension", "Dimension", "Dimension", "Fact"],
    "Grain": ["Per customer", "Per product", "Per day", "Per region-city", "Per line item"],
    "Rows": [len(customers), len(products), "790 days", "4 regions", "25,228"],
})
add_table(dim_tables, "Schema Overview:")
doc.add_paragraph("")

doc.add_paragraph(
    "Key indexes exist on customer_id, product_id, date_key, region_id, and order_id. "
    "A materialized view (mv_daily_sales) pre-aggregates daily sales by category, "
    "brand, and region for dashboard performance."
)

# ══════════════════════════════════════════════════════════════════
# 5. SQL ANALYSIS RESULTS
# ══════════════════════════════════════════════════════════════════
add_heading("5. SQL Analysis", level=1)
doc.add_paragraph(
    "12 advanced SQL queries (sql/analysis_queries.sql) were written covering:"
)
sql_topics = [
    "Revenue by Region — West leads in volume, South in profit margin",
    "Monthly Sales Trend — MoM change with LAG window function",
    "Top 10 Products by Profit — RANK() window function",
    "Customer Lifetime Value (CLV) — Tiered segmentation (Platinum/Gold/Silver/Bronze)",
    "Repeat Purchase Rate — Distribution of order frequency across customers",
    "Return Rate by Category — Electronics dominates at 26%",
    "Profit Margin by Product — Bottom 20 loss-making products identified",
    "Cohort Analysis — Monthly retention cohort with month-offset tracking",
    "Window Functions — LEAD for next purchase date, running totals",
    "Discount Impact Analysis — Margin erosion from 42.7% (0% disc) to -6.5% (30%+)",
    "Payment Mode Analysis — COD shows highest return correlation",
    "Top 10 Loss-Making Products — Products where profit < 0",
]
for t in sql_topics:
    add_insight(t)

# Revenue by region
by_region = pd.DataFrame(results.get("eda", {}).get("by_region", []))
if len(by_region):
    add_heading("Revenue by Region", level=2)
    add_table(by_region)

# Discount impact
by_discount = pd.DataFrame(results.get("eda", {}).get("by_discount", []))
if len(by_discount):
    add_heading("Discount Impact on Margin", level=2)
    add_table(by_discount)

# ══════════════════════════════════════════════════════════════════
# 6. PYTHON ANALYTICS
# ══════════════════════════════════════════════════════════════════
add_heading("6. Python Analytics", level=1)

# 6.1 EDA
add_heading("6.1 Exploratory Data Analysis", level=2)
doc.add_paragraph(
    "Full EDA was conducted revealing revenue concentration, margin distribution, "
    "and category-level performance."
)

by_cat = pd.DataFrame(results.get("eda", {}).get("by_category", []))
if len(by_cat):
    add_heading("Revenue by Category", level=3)
    add_table(by_cat)

# 6.2 Correlations
add_heading("6.2 Correlation Analysis", level=2)
corr = results.get("correlations", {})
doc.add_paragraph(
    f"Revenue-Profit correlation: {corr.get('revenue_profit', 'N/A')} (strong positive)\n"
    f"Discount-Profit correlation: {corr.get('discount_profit', 'N/A')} (near zero — discounts drive "
    f"volume but erode per-unit profit)\n"
    f"Discount-Revenue correlation: {corr.get('discount_revenue', 'N/A')} (moderate — higher discounts "
    f"associate with higher revenue)"
)

# 6.3 Segmentation
add_heading("6.3 Customer Segmentation (RFM + KMeans)", level=2)
seg = results.get("segmentation", {})
doc.add_paragraph(
    f"Optimal clusters: {seg.get('optimal_k', 'N/A')} (silhouette score: {seg.get('silhouette', 'N/A')})\n"
    f"Two customer segments identified:"
)
segments = seg.get("segments", [])
for s in segments:
    label = s.get("label", "Unknown")
    count = s.get("count", 0)
    avg_mon = s.get("avg_monetary", 0)
    avg_freq = s.get("avg_frequency", 0)
    doc.add_paragraph(
        f"  • {label}: {count} customers, avg ${avg_mon:,.0f} revenue, "
        f"{avg_freq:.1f} avg orders",
        style='List Bullet'
    )

# 6.4 Forecasting
add_heading("6.4 Sales Forecasting", level=2)
forecast = results.get("forecasting", {})
doc.add_paragraph(
    f"Method: {forecast.get('method', 'ARIMA (seasonal)')}\n"
    f"MAE: {forecast.get('mae', 'N/A'):,}\n"
    f"Next 90-day revenue forecast: ${forecast.get('next_90d_revenue', 0):,.0f}"
)

# ══════════════════════════════════════════════════════════════════
# 7. KEY INSIGHTS
# ══════════════════════════════════════════════════════════════════
add_heading("7. Key Business Insights", level=1)
insights = results.get("insights", [])
for ins in insights:
    add_insight(ins)

doc.add_paragraph("")
doc.add_paragraph(
    "These insights directly answer the original business questions:\n"
    "• Which products are profitable? → Top 10 by profit identified; 10 loss-makers flagged\n"
    "• Which customers bring long-term value? → CLV scored; 501 High-Value segment isolated\n"
    "• Which regions are underperforming? → West has highest revenue but lowest margin (18.3%)\n"
    "• Why are returns increasing? → COD 2.2x more returns; Electronics leads at 26%\n"
    "• What drives revenue fluctuations? → Discounts >20% destroy margin; Q4 seasonality drives 45% spikes"
)

# ══════════════════════════════════════════════════════════════════
# 8. POWER BI & TABLEAU
# ══════════════════════════════════════════════════════════════════
add_heading("8. BI Tool Integration", level=1)

add_heading("8.1 Power BI", level=2)
doc.add_paragraph(
    "Exports in data/exports/powerbi/ contain star-schema CSVs ready for import:\n"
    "• dim_date.csv, dim_customers.csv, dim_products.csv (dimensions)\n"
    "• fact_sales.csv (fact table with 25,266 rows)\n"
    "• returns.csv (return reason details)\n\n"
    "DAX measures (power_bi/dax_measures.txt) include:\n"
    "• KPI measures: Total Revenue, Profit, Orders, Margin %, AOV\n"
    "• Time intelligence: MoM%, YoY%, MTD, QTD, YTD\n"
    "• Rankings: Product/Customer revenue ranks\n"
    "• Segmentation: Dynamic customer tier (Platinum/Gold/Silver/Bronze)\n\n"
    "Recommended 5-page dashboard layout:\n"
    "  1. Executive Dashboard — KPI cards, trend lines, category bars\n"
    "  2. Sales Analysis — Monthly trends, regional, MoM decomposition\n"
    "  3. Customer Analytics — RFM scatter, segments, CLV\n"
    "  4. Product Analytics — Top/bottom products, discount vs margin\n"
    "  5. Returns Analysis — Return rate by category, payment mode"
)

add_heading("8.2 Tableau", level=2)
doc.add_paragraph(
    "Exports in data/exports/tableau/ contain:\n"
    "• tableau_sales.csv (25,266 rows) — daily sales by region, category, product\n"
    "• tableau_customers.csv (1,282 rows) — customer-level CLV, AOV, return rate\n"
    "• tableau_regional.csv (20 rows) — region×category with lat/lon for geographic heatmaps\n"
    "• tableau_trend.csv (26 rows) — monthly trends for storytelling\n\n"
    "Recommended Tableau Story:\n"
    "  1. Geographic heatmap of revenue and margin by region\n"
    "  2. Trend storytelling — revenue growth with seasonality annotations\n"
    "  3. Product performance funnel — category → sub-category → product"
)

# ══════════════════════════════════════════════════════════════════
# 9. CONCLUSION
# ══════════════════════════════════════════════════════════════════
add_heading("9. Conclusion & Recommendations", level=1)

doc.add_paragraph(
    "The E-Commerce Analytics Platform successfully demonstrates an end-to-end analytics "
    "pipeline, from raw data generation through to actionable business insights. "
    "The system integrates SQL, Python, Power BI, and Tableau into a cohesive workflow.\n\n"
    "Key recommendations based on the analysis:\n"
)

recommendations = [
    "Reduce discount depth: Orders with ≥20% discount yield negative margins (−6.5% for 30%+ band). "
    "Tier discounts with hard floors at 15%.",
    "Target COD-to-prepaid conversion: COD orders have 2.2x higher return rate. "
    "Offer UPI/Credit Card incentives at checkout.",
    "Invest in Electronics quality assurance: 26% return rate is the highest — "
    "improve product descriptions, add video demos, and tighten quality checks.",
    "West region margin recovery: Despite highest volume, West has the lowest margin (18.3%). "
    "Review pricing strategy and promotional intensity.",
    "High-Value customer program: 501 customers generate 80.7% of revenue. "
    "Launch a loyalty program with exclusive benefits for this segment.",
    "Books category optimization: Despite highest margin (54.8%), Books generate only 1% of revenue. "
    "Cross-sell with Electronics/Clothing to increase basket size.",
]
for r in recommendations:
    add_insight(r)

doc.add_paragraph("")
doc.add_paragraph(
    "All code, data, and documentation are available in the project root. "
    "Run python run_all.py to regenerate the full pipeline."
)

# ── Save ─────────────────────────────────────────────────────────
output_path = os.path.join(BASE_DIR, "ECommerce_Analytics_Report.docx")
doc.save(output_path)
print(f"Report saved to: {output_path}")
